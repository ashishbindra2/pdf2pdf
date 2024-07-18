from pdf2docx import Converter
import os
import yaml
from docx import Document
from googletrans import Translator
# from langdetect import detect
from pydantic import BaseModel

# Load configuration
with open('config.yaml') as config_file:
    config = yaml.safe_load(config_file)

input_pdf = config['paths'].get('pdf_path')
output_docx = config['paths'].get('docx_path')

if not input_pdf or not output_docx:
    raise ValueError("Configuration paths for 'pdf_path' and 'docx_path' must be set in config.yaml")


# Convert pdf to docx
def convert_pdf2docx(pdf_file: str):
    docx_file = pdf_file.replace(".pdf", ".docx")

    pdf_file_input = os.path.join(input_pdf, pdf_file)
    pdf_file_output = os.path.join(output_docx, docx_file)

    # Normalize the paths for the current operating system
    pdf_file_input = os.path.normpath(pdf_file_input)
    pdf_file_output = os.path.normpath(pdf_file_output)

    print(f"input {pdf_file_input}")

    if os.path.isfile(pdf_file_input):
        cv = Converter(pdf_file_input)
        cv.convert(pdf_file_output)  # all pages by default
        cv.close()
    else:
        raise FileNotFoundError(f"File not found: {pdf_file_input}")

    return pdf_file_output


def translate_text(docx: str, target_language="nl"):
    # Initialize the translator
    translator = Translator()

    doc = Document(docx)

    # Load the DOCX file /mnt/data/
    file_name = docx.split("\\")[-1]
    print(f"{file_name=:}", target_language)

    # Translate paragraphs
    for para in doc.paragraphs:
        if para.text.strip():  # Check if the paragraph is not empty
            print(para.text)
            # print(translator.detect(para.text))
            translated_text = translator.translate(para.text, dest=target_language).text
            para.text = translated_text

    # Translate text in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():  # Check if the cell is not empty
                    try:
                        print(translator.detect(cell.text))

                        translated_text = translator.translate(cell.text, dest=target_language)
                        if translated_text is None:
                            continue
                        translated_text = translated_text.text
                        cell.text = translated_text
                    except Exception as e:
                        print("yo error")

    # Save the modified document
    docx_file = os.path.join(output_docx, f'Modified_{target_language.capitalize()}_{file_name}')

    docx_file = os.path.normpath(docx_file)
    print(f"{docx_file=:}")
    doc.save(docx_file)
    del doc
    return docx_file
